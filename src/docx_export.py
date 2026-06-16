"""Генерация DOCX-документа программы ДПО из Markdown.

Скрипт читает Markdown-файл программы и генерирует DOCX с форматированием
по ГОСТу: Times New Roman 14pt, поля, таблицы, нумерация разделов.

Использование:
    python docx_export.py [--input output/program_dpo.md] [--output output/program_dpo.docx]
"""

import re
import sys
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def configure_section(section, orientation=WD_ORIENT.PORTRAIT):
    """Настроить секцию документа по ГОСТу."""
    section.orientation = orientation
    if orientation == WD_ORIENT.PORTRAIT:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    else:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)


def set_cell_font(cell, text, size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font_name="Times New Roman"):
    """Установить текст ячейки таблицы с форматированием."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    # Для кириллицы
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text, bold=False, size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    """Добавить параграф с форматированием."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), "Times New Roman")
    rFonts.set(qn('w:cs'), "Times New Roman")
    rPr.insert(0, rFonts)
    return p


def merge_cells(table, row, start_col, end_col):
    """Объединить ячейки в строке таблицы."""
    cell_start = table.cell(row, start_col)
    cell_end = table.cell(row, end_col)
    cell_start.merge(cell_end)


def add_table_borders(table):
    """Добавить границы таблице."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


def parse_markdown(md_text):
    """Разобрать Markdown в структурированные блоки."""
    blocks = []
    current_table = []
    in_table = False

    for line in md_text.split('\n'):
        stripped = line.strip()

        # Таблица
        if stripped.startswith('|') and '|' in stripped[1:]:
            if '---' in stripped:
                continue  # Skip separator row
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                current_table = []
            current_table.append(cells)
            continue
        else:
            if in_table:
                blocks.append({'type': 'table', 'rows': current_table})
                current_table = []
                in_table = False

        # Заголовки
        if stripped.startswith('# '):
            blocks.append({'type': 'h1', 'text': stripped[2:]})
        elif stripped.startswith('## '):
            blocks.append({'type': 'h2', 'text': stripped[3:]})
        elif stripped.startswith('### '):
            blocks.append({'type': 'h3', 'text': stripped[4:]})
        elif stripped.startswith('#### '):
            blocks.append({'type': 'h4', 'text': stripped[5:]})
        elif stripped.startswith('- ') or stripped.startswith('* '):
            blocks.append({'type': 'list_item', 'text': stripped[2:]})
        elif stripped.startswith('**') and stripped.endswith('**'):
            blocks.append({'type': 'bold_text', 'text': stripped[2:-2]})
        elif stripped == '':
            blocks.append({'type': 'empty'})
        else:
            # Обработка **bold** внутри текста
            blocks.append({'type': 'paragraph', 'text': stripped})

    if in_table:
        blocks.append({'type': 'table', 'rows': current_table})

    return blocks


def md_to_docx(md_path, docx_path):
    """Конвертировать Markdown-файл программы ДПО в DOCX."""
    md_text = Path(md_path).read_text(encoding='utf-8')
    blocks = parse_markdown(md_text)

    doc = Document()
    configure_section(doc.sections[0])

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    i = 0
    while i < len(blocks):
        block = blocks[i]

        if block['type'] == 'h1':
            add_paragraph(doc, block['text'], bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

        elif block['type'] == 'h2':
            add_paragraph(doc, block['text'], bold=True, size=14, space_after=8)

        elif block['type'] == 'h3':
            add_paragraph(doc, block['text'], bold=True, size=14, space_after=6)

        elif block['type'] == 'h4':
            add_paragraph(doc, block['text'], bold=True, size=14, space_after=4)

        elif block['type'] == 'bold_text':
            add_paragraph(doc, block['text'], bold=True, space_after=4)

        elif block['type'] == 'list_item':
            add_paragraph(doc, f"• {block['text']}", space_after=2)

        elif block['type'] == 'table':
            rows_data = block['rows']
            if not rows_data:
                i += 1
                continue

            num_cols = len(rows_data[0])
            num_rows = len(rows_data)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            add_table_borders(table)

            for row_idx, row_data in enumerate(rows_data):
                for col_idx, cell_text in enumerate(row_data[:num_cols]):
                    cell = table.cell(row_idx, col_idx)
                    is_header = (row_idx == 0)
                    set_cell_font(
                        cell, cell_text,
                        size=11 if num_cols > 4 else 14,
                        bold=is_header,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                    )

            doc.add_paragraph()  # Spacer after table

        elif block['type'] == 'paragraph':
            # Handle inline bold
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)

            # Split by ** for bold segments
            parts = re.split(r'(\*\*.*?\*\*)', block['text'])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)

                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), "Times New Roman")
                rFonts.set(qn('w:cs'), "Times New Roman")
                rPr.insert(0, rFonts)

        elif block['type'] == 'empty':
            pass  # Skip empty lines

        i += 1

    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")
    return docx_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Конвертер программы ДПО из Markdown в DOCX")
    parser.add_argument("--input", "-i", default="output/program_dpo.md", help="Путь к входному Markdown-файлу")
    parser.add_argument("--output", "-o", default="output/program_dpo.docx", help="Путь к выходному DOCX-файлу")
    args = parser.parse_args()

    base_dir = Path(__file__).parent.parent
    input_path = base_dir / args.input
    output_path = base_dir / args.output

    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        print(f"❌ Файл не найден: {input_path}")
        sys.exit(1)

    md_to_docx(input_path, output_path)